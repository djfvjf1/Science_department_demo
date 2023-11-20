import calendar
import datetime
import os
from io import BytesIO

from django.conf import settings
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import UserCreationForm
from django.contrib.auth.mixins import LoginRequiredMixin
from django.http import HttpResponse, HttpResponseRedirect
from django.shortcuts import redirect, render
from django.urls import reverse_lazy
from django.views.generic.edit import CreateView
from docx import Document

from .forms import (KPI_1_Form, KPI_2_Form, KPI_3_Form, KPI_4_Form, KPI_5_Form,
                    KPI_6_Form, KPI_7_Form, KPI_8_Form, KPI_9_Form)
from .models import (KPI_1, KPI_2, KPI_3, KPI_4, KPI_5, KPI_6, KPI_7, KPI_8,
                     KPI_9, TeachersAddInfo)


@login_required
def home(request):
    user_id = request.user.id
    try:
        user_info = TeachersAddInfo.objects.get(user_id=user_id)
    except TeachersAddInfo.DoesNotExist:
        user_info = None

    tables = KPI_1.objects.filter(author=user_info.user_id)
    tables2 = KPI_2.objects.filter(author=user_info.user_id)
    tables3 = KPI_3.objects.filter(author=user_info.user_id)
    tables4 = KPI_4.objects.filter(author=user_info.user_id)
    tables5 = KPI_5.objects.filter(author=user_info.user_id)
    tables6 = KPI_6.objects.filter(author=user_info.user_id)
    tables7 = KPI_7.objects.filter(author=user_info.user_id)
    tables8 = KPI_8.objects.filter(author=user_info.user_id)
    tables9 = KPI_9.objects.filter(author=user_info.user_id)
    
    return render(request, "home/home.html", {'user_info': user_info, 'user': request.user,
                                "tables": tables, 
                                "tables2": tables2, 
                                "tables3": tables3,
                                "tables4": tables4,
                                "tables5": tables5,
                                "tables6": tables6,
                                "tables7": tables7,
                                "tables8": tables8,
                                "tables9": tables9
                                
      })

@login_required
def kpi(request):
    return render(request, "home/KPI.html", {'kpi': kpi})

def prehome(request):
    return render(request, "home/prehome.html", {"prehome": prehome})

# Таблица №20
@login_required
def new_form(request):
    
    if request.method == 'POST':
        form = KPI_1_Form(request.POST, request.FILES)
        # check if form data is valid
        if form.is_valid():
            #form.save(commit=False)
            form.author_name = request.user.username
            table = form.save()
            return redirect('/table/' + str(table.id) + '/')
    else:
        form = KPI_1_Form()
     

    return render(request, "form/form.html", {'form': form})

@login_required
def table_page(request, pk):
    table = KPI_1.objects.get(id=pk)
    table.refresh_from_db()
    return render(request, "form/table_page.html", {'table': table})

@login_required
def edit_form(request, pk):
    table = KPI_1.objects.get(id=pk)
    
    if request.method == 'POST':
        form = KPI_1_Form(request.POST, request.FILES, instance=table)
        if form.is_valid():
            form.author_name = request.user.username
            # Update the uploaded file if a new file is provided
            if 'uploaded_file' in request.FILES:
                table.uploaded_file = request.FILES['uploaded_file']
            # Save the form and update the instance
            form.save()
            return redirect('/table/' + str(table.id) + '/')
    else:
        form = KPI_1_Form(instance=table)
     
    return render(request, "form/edit_form.html", {'form': form})

# Таблица №2
@login_required
def new_form2(request):
    if request.method == 'POST':
        form = KPI_2_Form(request.POST, request.FILES)
        if form.is_valid():
            kpi_2_instance = form.save(commit=False)
            kpi_2_instance.subject = request.user
            kpi_2_instance.save()
            return redirect('/table2/' + str(kpi_2_instance.id) + '/')
    else:
        form = KPI_2_Form()

    return render(request, "table2/form.html", {'form': form})

# Update the table_page2 view for KPI_2
@login_required
def table_page2(request, pk):
    table = KPI_2.objects.get(id=pk)
    table.refresh_from_db()
    return render(request, "table2/table_page2.html", {'table': table})

@login_required
def edit_form2(request, pk):
    table = KPI_2.objects.get(id=pk)
    
    if request.method == 'POST':
        form = KPI_2_Form(request.POST, request.FILES, instance=table)
        if form.is_valid():
            form.title_of_the_monograph = request.user.username
            # Update the uploaded file if a new file is provided
            if 'uploaded_file' in request.FILES:
                table.uploaded_file = request.FILES['uploaded_file']
            # Save the form and update the instance
            form.save()
            return redirect('/table2/' + str(table.id) + '/')
    else:
        form = KPI_2_Form(instance=table)
     
    return render(request, "table2/edit_form2.html", {'form': form})


# Update the new_form3 view for KPI_3
@login_required
def new_form3(request):
    if request.method == 'POST':
        form = KPI_3_Form(request.POST, request.FILES)
        if form.is_valid():
            kpi_3_instance = form.save(commit=False)
            kpi_3_instance.organisation = request.user
            kpi_3_instance.save()
            return redirect('/table3/' + str(kpi_3_instance.id) + '/')
    else:
        form = KPI_3_Form()

    return render(request, "table3/form.html", {'form': form})

# Update the table_page3 view for KPI_3
@login_required
def table_page3(request, pk):
    table = KPI_3.objects.get(id=pk)
    table.refresh_from_db()
    return render(request, "table3/table_page3.html", {'table': table})

@login_required
def edit_form3(request, pk):
    table = KPI_3.objects.get(id=pk)
    
    if request.method == 'POST':
        form = KPI_3_Form(request.POST, request.FILES, instance=table)
        if form.is_valid():
            form.title_of_the_project = request.user.username
            # Update the uploaded file if a new file is provided
            if 'uploaded_file' in request.FILES:
                table.uploaded_file = request.FILES['uploaded_file']
            # Save the form and update the instance
            form.save()
            return redirect('/table3/' + str(table.id) + '/')
    else:
        form = KPI_3_Form(instance=table)
     
    return render(request, "table3/edit_form3.html", {'form': form})


# Таблица №4
@login_required
def new_form4(request):
    
    if request.method == 'POST':
        form = KPI_4_Form(request.POST, request.FILES)
        # check if form data is valid
        if form.is_valid():
            #form.save(commit=False)
            form.event_title = request.user
            table = form.save()
            return redirect('/table4/' + str(table.id) + '/')
    else:
        form = KPI_4_Form()
     

    return render(request, "table4/form.html", {'form': form})

@login_required
def table_page4(request, pk):
    table = KPI_4.objects.get(id=pk)
    table.refresh_from_db()
    return render(request, "table4/table_page4.html", {'table': table})

@login_required
def edit_form4(request, pk):
    table = KPI_4.objects.get(id=pk)
    
    if request.method == 'POST':
        form = KPI_4_Form(request.POST, request.FILES, instance=table)
        if form.is_valid():
            form.event_title = request.user.username
            # Update the uploaded file if a new file is provided
            if 'uploaded_file' in request.FILES:
                table.uploaded_file = request.FILES['uploaded_file']
            # Save the form and update the instance
            form.save()
            return redirect('/table4/' + str(table.id) + '/')
    else:
        form = KPI_4_Form(instance=table)
     
    return render(request, "table4/edit_form4.html", {'form': form})

# Таблица №5
@login_required
def new_form5(request):
    
    if request.method == 'POST':
        form = KPI_5_Form(request.POST, request.FILES)
        # check if form data is valid
        if form.is_valid():
            #form.save(commit=False)
            form.full_name_of_the_doctoral_student = request.user
            table = form.save()
            return redirect('/table5/' + str(table.id) + '/')
    else:
        form = KPI_5_Form()
     

    return render(request, "table5/form.html", {'form': form})

@login_required
def table_page5(request, pk):
    table = KPI_5.objects.get(id=pk)
    table.refresh_from_db()
    return render(request, "table5/table_page5.html", {'table': table})

@login_required
def edit_form5(request, pk):
    table = KPI_5.objects.get(id=pk)
    
    if request.method == 'POST':
        form = KPI_5_Form(request.POST, request.FILES, instance=table)
        if form.is_valid():
            form.full_name_of_the_doctoral_student = request.user.username
            # Update the uploaded file if a new file is provided
            if 'uploaded_file' in request.FILES:
                table.uploaded_file = request.FILES['uploaded_file']
            # Save the form and update the instance
            form.save()
            return redirect('/table5/' + str(table.id) + '/')
    else:
        form = KPI_4_Form(instance=table)
     
    return render(request, "table5/edit_form5.html", {'form': form})

# Таблица №6
@login_required
def new_form6(request):
    
    if request.method == 'POST':
        form = KPI_6_Form(request.POST, request.FILES)
        # check if form data is valid
        if form.is_valid():
            #form.save(commit=False)
            form.full_name_of_the_doctoral_student = request.user
            table = form.save()
            return redirect('/table6/' + str(table.id) + '/')
    else:
        form = KPI_6_Form()
     

    return render(request, "table6/form.html", {'form': form})

@login_required
def table_page6(request, pk):
    table = KPI_6.objects.get(id=pk)
    table.refresh_from_db()
    return render(request, "table6/table_page6.html", {'table': table})

@login_required
def edit_form6(request, pk):
    table = KPI_6.objects.get(id=pk)
    
    if request.method == 'POST':
        form = KPI_6_Form(request.POST, request.FILES, instance=table)
        if form.is_valid():
            form.full_name_of_the_doctoral_student = request.user.username
            # Update the uploaded file if a new file is provided
            if 'uploaded_file' in request.FILES:
                table.uploaded_file = request.FILES['uploaded_file']
            # Save the form and update the instance
            form.save()
            return redirect('/table6/' + str(table.id) + '/')
    else:
        form = KPI_6_Form(instance=table)
     
    return render(request, "table6/edit_form6.html", {'form': form})

# Таблица №7
@login_required
def new_form7(request):
    
    if request.method == 'POST':
        form = KPI_7_Form(request.POST, request.FILES)
        # check if form data is valid
        if form.is_valid():
            #form.save(commit=False)
            form.decision_of_the_authorized_body_on_awarding_an_academic_degree_in_the_field_of_art_or_sports = request.user
            table = form.save()
            return redirect('/table7/' + str(table.id) + '/')
    else:
        form = KPI_7_Form()
     

    return render(request, "table7/form.html", {'form': form})

@login_required
def table_page7(request, pk):
    table = KPI_7.objects.get(id=pk)
    table.refresh_from_db()
    return render(request, "table7/table_page7.html", {'table': table})

@login_required
def edit_form7(request, pk):
    table = KPI_7.objects.get(id=pk)
    
    if request.method == 'POST':
        form = KPI_7_Form(request.POST, request.FILES, instance=table)
        if form.is_valid():
            form.decision_of_the_authorized_body_on_awarding_an_academic_degree_in_the_field_of_art_or_sports = request.user.username
            # Update the uploaded file if a new file is provided
            if 'uploaded_file' in request.FILES:
                table.uploaded_file = request.FILES['uploaded_file']
            # Save the form and update the instance
            form.save()
            return redirect('/table7/' + str(table.id) + '/')
    else:
        form = KPI_7_Form(instance=table)
     
    return render(request, "table7/edit_form7.html", {'form': form})

# Таблица №8
@login_required
def new_form8(request):
    
    if request.method == 'POST':
        form = KPI_8_Form(request.POST, request.FILES)
        # check if form data is valid
        if form.is_valid():
            #form.save(commit=False)
            form.decision_of_the_AlmaU_School_or_management_on_the_scientific_management_of_the_laboratory_or_center = request.user
            table = form.save()
            return redirect('/table8/' + str(table.id) + '/')
    else:
        form = KPI_8_Form()
     

    return render(request, "table8/form.html", {'form': form})

@login_required
def table_page8(request, pk):
    table = KPI_8.objects.get(id=pk)
    table.refresh_from_db()
    return render(request, "table8/table_page8.html", {'table': table})

@login_required
def edit_form8(request, pk):
    table = KPI_8.objects.get(id=pk)
    
    if request.method == 'POST':
        form = KPI_8_Form(request.POST, request.FILES, instance=table)
        if form.is_valid():
            form.decision_of_the_AlmaU_School_or_management_on_the_scientific_management_of_the_laboratory_or_center = request.user.username
            # Update the uploaded file if a new file is provided
            if 'uploaded_file' in request.FILES:
                table.uploaded_file = request.FILES['uploaded_file']
            # Save the form and update the instance
            form.save()
            return redirect('/table8/' + str(table.id) + '/')
    else:
        form = KPI_8_Form(instance=table)
     
    return render(request, "table8/edit_form8.html", {'form': form})

# Таблица №9
@login_required
def new_form9(request):
    
    if request.method == 'POST':
        form = KPI_9_Form(request.POST, request.FILES)
        # check if form data is valid
        if form.is_valid():
            #form.save(commit=False)
            form.name_of_the_patent = request.user
            table = form.save()
            return redirect('/table9/' + str(table.id) + '/')
    else:
        form = KPI_9_Form()
     

    return render(request, "table9/form.html", {'form': form})

@login_required
def table_page9(request, pk):
    table = KPI_9.objects.get(id=pk)
    table.refresh_from_db()
    return render(request, "table9/table_page9.html", {'table': table})

@login_required
def edit_form9(request, pk):
    table = KPI_9.objects.get(id=pk)
    
    if request.method == 'POST':
        form = KPI_9_Form(request.POST, request.FILES, instance=table)
        if form.is_valid():
            form.name_of_the_patent = request.user.username
            # Update the uploaded file if a new file is provided
            if 'uploaded_file' in request.FILES:
                table.uploaded_file = request.FILES['uploaded_file']
            # Save the form and update the instance
            form.save()
            return redirect('/table9/' + str(table.id) + '/')
    else:
        form = KPI_9_Form(instance=table)
     
    return render(request, "table9/edit_form9.html", {'form': form})

@login_required
def export_data_to_word_for_table(request):
    # Get your queryset of data
    queryset = KPI_1.objects.all()

    # Create a new Word document
    document = Document()

    # Iterate over the data and group it by month
    data_by_month = {}
    for obj in queryset:
        month = obj.date.month
        if month not in data_by_month:
            data_by_month[month] = []
        data_by_month[month].append(obj)

    # Add a table for each month's data
    for month, data in data_by_month.items():
        # Add a section heading with the month
        document.add_heading(calendar.month_name[month], level=1)

        # Add a table to the document with headers
        table = document.add_table(rows=2, cols=5)  # Added one more column for the file link
        hdr_cells = table.rows[0].cells
        table.style = 'Table Grid'
        hdr_cells[0].text = 'Ф.И.О. автора'
        hdr_cells[1].text = 'Название'
        hdr_cells[2].text = 'Название журнала'
        hdr_cells[3].text = 'Год публикации'
        hdr_cells[4].text = 'Ссылка на файл'

        # Add data rows to the table
        for obj in data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(obj.author_name)
            row_cells[1].text = str(obj.title_of_the_article)
            row_cells[2].text = str(obj.name_of_the_magazine)
            row_cells[3].text = str(obj.year_of_publication)

            # Add a link to the uploaded file
            file_url = os.path.join(settings.MEDIA_URL, str(obj.uploaded_file))
            row_cells[4].text = file_url

    # Set the file name and Content-Type header for the response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=exported_data.docx'

    # Save the Word document to the response
    document.save(response)
    return response

#------------------------------------------------------------------------------------------
@login_required
def export_data_to_word_for_table2(request):
    # Get your queryset of data
    queryset = KPI_2.objects.all()

    # Create a new Word document
    document = Document()

    # Iterate over the data and group it by month
    data_by_month2 = {}
    for obj in queryset:
        month = obj.date.month
        if month not in data_by_month2:
            data_by_month2[month] = []
        data_by_month2[month].append(obj)

    # Add a table for each month's data
    for month, data in data_by_month2.items():
        # Add a section heading with the month
        document.add_heading(calendar.month_name[month], level=1)

        # Add a table to the document with headers
        table = document.add_table(rows=2, cols=5)  # Added one more column for the file link
        hdr_cells = table.rows[0].cells
        table.style = 'Table Grid'
        hdr_cells[0].text = 'Название монографии'
        hdr_cells[1].text = 'Год издательства'
        hdr_cells[2].text = 'Название издательства'
        hdr_cells[3].text = 'Количество печатных листов'


        # Add data rows to the table
        for obj in data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(obj.title_of_the_monograph)
            row_cells[1].text = str(obj.year_of_publication)
            row_cells[2].text = str(obj.name_of_publication)
            row_cells[3].text = str(obj.number_of_printed_sheets)

            # Add a link to the uploaded file
            file_url = os.path.join(settings.MEDIA_URL, str(obj.uploaded_file))
            row_cells[4].text = file_url

    # Set the file name and Content-Type header for the response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=exported_data.docx'

    # Save the Word document to the response
    document.save(response)
    return response

    


#------------------------------------------------------------------------------------------
@login_required
def export_data_to_word_for_table3(request):
    # Get your queryset of data
    queryset = KPI_2.objects.all()

    # Create a new Word document
    document = Document()

    # Iterate over the data and group it by month
    data_by_month2 = {}
    for obj in queryset:
        month = obj.date.month
        if month not in data_by_month2:
            data_by_month2[month] = []
        data_by_month2[month].append(obj)

    # Add a table for each month's data
    for month, data in data_by_month2.items():
        # Add a section heading with the month
        document.add_heading(calendar.month_name[month], level=1)

        # Add a table to the document with headers
        table = document.add_table(rows=2, cols=5)  # Added one more column for the file link
        hdr_cells = table.rows[0].cells
        table.style = 'Table Grid'
        hdr_cells[0].text = 'Название проекта'
        hdr_cells[1].text = 'Номер проекта'
        hdr_cells[2].text = 'Годы выполнения'
        hdr_cells[3].text = 'Роль в проекте'
        hdr_cells[4].text = 'Руководитель проекта'


        # Add data rows to the table
        for obj in data:
            row_cells = table.add_row().cells
            row_cells[0].text = str(obj.title_of_the_project)
            row_cells[1].text = str(obj.number_of_the_project)
            row_cells[2].text = str(obj.years_of_completion)
            row_cells[3].text = str(obj.role_in_the_project)
            row_cells[4].text = str(obj.director_of_the_project)


            # Add a link to the uploaded file
            file_url = os.path.join(settings.MEDIA_URL, str(obj.uploaded_file))
            row_cells[5].text = file_url

    # Set the file name and Content-Type header for the response
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=exported_data.docx'

    # Save the Word document to the response
    document.save(response)
    return response
    

#------------------------------------------------------------------------------------------

# @login_required
# def export_data_to_word(request):
    

#     # Create a new Word document
#     document = Document()



#     # Таблица №20

#     # Get your queryset of data

#     document.add_paragraph('Перечень патентов, полученных в 20____-20___ годы', style='Heading 1')

#     queryset = Table.objects.all()

#     # Iterate over the data and group it by month 
#     data_by_month20 = {} 
#     for obj in queryset: 
#         month = obj.date.month 
#         if month not in data_by_month20: 
#             data_by_month20[month] = [] 
#         data_by_month20[month].append(obj) 
 
#     # Add a table for each month's data 
#     for month, data in data_by_month20.items(): 
#         # Add a section heading with the month 
#         document.add_heading(calendar.month_name[month], level=1) 
 
#         # Add a table to the document with headers 
#         table = document.add_table(rows=2, cols=4)
#         hdr_cells = table.rows[0].cells
#         table.style = 'Table Grid'
#         hdr_cells[0].text = 'Ф.И.О. автора'
#         hdr_cells[1].text = '№ патента'
#         hdr_cells[2].text = 'Год выдачи'
#         hdr_cells[3].text = 'Название'
 
#         # Add data rows to the table 
#         for obj in data: 
#             row_cells = table.add_row().cells 
#             row_cells[0].text = str(obj.author_name) 
#             row_cells[1].text = str(obj.patent) 
#             row_cells[2].text = str(obj.year) 
#             row_cells[3].text = str(obj.title) 
 

#     # Таблица №23
#     document.add_paragraph('Договора о международном сотрудничестве', style='Heading 1')
#     # Get your queryset of data
#     queryset23 = Table23.objects.all()

#         # Iterate over the data and group it by month 
#     data_by_month23 = {} 
#     for obj in queryset23: 
#         month = obj.date.month 
#         if month not in data_by_month23: 
#             data_by_month23[month] = [] 
#         data_by_month23[month].append(obj) 
 
#     # Add a table for each month's data 
#     for month, data in data_by_month23.items(): 
#         # Add a section heading with the month 
#         document.add_heading(calendar.month_name[month], level=1) 
 
#         # Add a table to the document with headers 
#         table23 = document.add_table(rows=2, cols=6)
#         hdr_cells23 = table23.rows[0].cells
#         table23.style = 'Table Grid'
#         hdr_cells23[0].text = 'Предмет международного договора или проекта'
#         hdr_cells23[1].text = 'Наименование учебного заведения-партнера'
#         hdr_cells23[2].text = 'Дата заключения проектов и/илидоговоров'
#         hdr_cells23[3].text = 'дата начала'
#         hdr_cells23[4].text = 'дата окончания'
#         hdr_cells23[5].text = 'Фактическое наличие на момент проф. контроля'
 
#         # Add data rows to the table 
#         for obj in data: 
#             row_cells23 = table23.add_row().cells
#             row_cells23[0].text = str(obj.subject)
#             row_cells23[1].text = str(obj.name_of_partner)
#             date_str = obj.start_date.strftime("%Y-%m-%d")
#             date_str1 = obj.date_of_contract.strftime("%Y-%m-%d")
#             date_str2 = obj.end_date.strftime("%Y-%m-%d")
#             row_cells23[2].text = date_str1
#             row_cells23[3].text = date_str
#             row_cells23[4].text = date_str2
#             row_cells23[5].text = str(obj.availability)

    
#     # Таблица №26
#     document.add_paragraph('Соглашения о сотрудничестве с организациями образования или научными или научно-образовательными или научно-производственными центрами (магистратура, докторантура)', style='Heading 1')
#     # Get your queryset of data
#     queryset26 = Table26.objects.all()

#    # Iterate over the data and group it by month 
#     data_by_month26 = {} 
#     for obj in queryset26: 
#         month = obj.date.month 
#         if month not in data_by_month26: 
#             data_by_month26[month] = [] 
#         data_by_month26[month].append(obj) 
 
#     # Add a table for each month's data 
#     for month, data in data_by_month26.items(): 
#         # Add a section heading with the month 
#         document.add_heading(calendar.month_name[month], level=1) 
 
#         table26 = document.add_table(rows=2, cols=5)
#         hdr_cells26 = table26.rows[0].cells
#         table26.style = 'Table Grid'
#         hdr_cells26[0].text = 'Организация'
#         hdr_cells26[1].text = 'Предмет Соглашения или договора'
#         hdr_cells26[2].text = 'На какие направления подготовки/специальности распространяется действие договора'
#         hdr_cells26[3].text = 'Дата подписания'
#         hdr_cells26[4].text = 'Сроки действия договора'
 
#         # Add data rows to the table 
#         for obj in data: 
#             row_cells26 = table26.add_row().cells
#             row_cells26[0].text = str(obj.organisation)
#             row_cells26[1].text = str(obj.subject_of_contract)
#             row_cells26[2].text = str(obj.directon_of_speciality)
#             date_str = obj.date_of_conclusion_of_the_contract.strftime("%Y-%m-%d")
#             row_cells26[3].text = str(obj.date_of_conclusion_of_the_contract)
#             row_cells26[4].text = str(obj.terms_of_the_contract)

       

#     # Set the file name and Content-Type header for the response
#     response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
#     response['Content-Disposition'] = 'attachment; filename=exported_data.docx'

#     # Save the Word document to the response
#     document.save(response)
#     return response


def login_view(request):
    if request.method == 'POST':
        username = request.POST['username']
        password = request.POST['password']
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('home')
        else:
            return render(request, 'login\login.html', {'error_message': 'Invalid login'})
    else:
        return render(request, 'login\login.html')

@login_required
def profile_view(request):
    user_id = request.user.id
    try:
        user_info = TeachersAddInfo.objects.get(user_id=user_id)
    except TeachersAddInfo.DoesNotExist:
        user_info = None

    #tables = KPI_1.objects.filter(author=user_info.user_id)
    #tables2 = KPI_2.objects.filter(author=user_info.user_id)
    #tables3 = KPI_3.objects.filter(author=user_info.user_id)
    #tables4 = KPI_4.objects.filter(author=user_info.user_id)
    #tables5 = KPI_5.objects.filter(author=user_info.user_id)
    #tables6 = KPI_6.objects.filter(author=user_info.user_id)
    #tables7 = KPI_7.objects.filter(author=user_info.user_id)
    #tables8 = KPI_8.objects.filter(author=user_info.user_id)
    #tables9 = KPI_9.objects.filter(author=user_info.user_id)

    return render(request, 'login/profile.html', {'user_info': user_info, 'user': request.user
            #"tables": tables, 
            #"tables2": tables2, 
            #"tables3": tables3,
            #"tables4": tables4,
            #"tables5": tables5,
            #"tables6": tables6,
            #"tables7": tables7,
            #"tables8": tables8,
            #"tables9": tables9
    })
    
@login_required
def logout_view(request):
    logout(request)
    return redirect('prehome')